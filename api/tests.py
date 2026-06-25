"""Regression tests for account-scoped template generation API.

These tests are intentionally focused on security boundaries and idempotency,
because those are the highest-risk areas when n8n retries failed workflows.
"""

from pathlib import Path
import tempfile

from django.contrib.auth import get_user_model
from django.core.files.base import ContentFile
from django.test import override_settings
from docx import Document
from rest_framework import status
from rest_framework.test import APITestCase

from .models import DocumentTemplate, GeneratedDocument, ServiceAPIKey

User = get_user_model()


def build_docx_template_bytes() -> bytes:
    """Create a small DOCX template containing docxtpl variables."""
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp_path = Path(tmp.name)

    try:
        document = Document()
        document.add_paragraph("سلام {{ recipient_name }}")
        document.add_paragraph("موضوع: {{ subject }}")
        document.save(tmp_path)
        return tmp_path.read_bytes()
    finally:
        tmp_path.unlink(missing_ok=True)


TEST_STORAGE = {
    "default": {"BACKEND": "django.core.files.storage.FileSystemStorage"},
    "staticfiles": {"BACKEND": "django.contrib.staticfiles.storage.StaticFilesStorage"},
}


@override_settings(STORAGES=TEST_STORAGE, MEDIA_ROOT=tempfile.mkdtemp())
class LetterGenerationAPITests(APITestCase):
    def setUp(self):
        self.user = User.objects.create_user(username="n8n-account", password="pass12345")
        self.other_user = User.objects.create_user(username="other-account", password="pass12345")
        self.staff = User.objects.create_user(username="staff", password="pass12345", is_staff=True)

        self.template = DocumentTemplate.objects.create(
            owner=self.user,
            name="نامه معرفی",
            key="introduction-letter",
            required_fields=["recipient_name", "subject"],
            sample_variables={"recipient_name": "شرکت نمونه", "subject": "معرفی"},
            is_active=True,
        )
        self.template.template_file.save(
            "intro.docx",
            ContentFile(build_docx_template_bytes()),
            save=True,
        )

        self.other_template = DocumentTemplate.objects.create(
            owner=self.other_user,
            name="قالب حساب دیگر",
            key="private-template",
            required_fields=[],
            is_active=True,
        )
        self.other_template.template_file.save(
            "other.docx",
            ContentFile(build_docx_template_bytes()),
            save=True,
        )

    def test_templates_are_scoped_to_authenticated_user(self):
        self.client.force_authenticate(self.user)
        response = self.client.get("/api/templates/")

        self.assertEqual(response.status_code, status.HTTP_200_OK)
        keys = [item["key"] for item in response.data["templates"]]
        self.assertEqual(keys, ["introduction-letter"])

    def test_unauthenticated_requests_are_rejected(self):
        response = self.client.get("/api/templates/")
        self.assertEqual(response.status_code, status.HTTP_401_UNAUTHORIZED)

    def test_generate_document_returns_download_url_and_audit_fields(self):
        self.client.force_authenticate(self.user)
        response = self.client.post(
            "/api/generate/",
            {
                "template_key": "introduction-letter",
                "variables": {
                    "recipient_name": "شرکت نمونه",
                    "subject": "معرفی جهت همکاری",
                },
                "meta": {"workflow": "letter-generator"},
                "idempotency_key": "execution-1001",
                "source": "n8n",
            },
            format="json",
        )

        self.assertEqual(response.status_code, status.HTTP_201_CREATED)
        payload = response.data["document"]
        self.assertEqual(payload["status"], GeneratedDocument.Status.SUCCESS)
        self.assertEqual(payload["owner"]["username"], "n8n-account")
        self.assertEqual(payload["created_by"]["username"], "n8n-account")
        self.assertTrue(payload["download_url"])
        self.assertTrue(payload["storage_key"].endswith(".docx"))

    def test_idempotency_key_reuses_existing_document(self):
        self.client.force_authenticate(self.user)
        body = {
            "template_key": "introduction-letter",
            "variables": {"recipient_name": "الف", "subject": "ب"},
            "idempotency_key": "same-execution-id",
            "source": "n8n",
        }

        first = self.client.post("/api/generate/", body, format="json")
        second = self.client.post("/api/generate/", body, format="json")

        self.assertEqual(first.status_code, status.HTTP_201_CREATED)
        self.assertEqual(second.status_code, status.HTTP_200_OK)
        self.assertFalse(first.data["reused"])
        self.assertTrue(second.data["reused"])
        self.assertEqual(first.data["document"]["id"], second.data["document"]["id"])
        self.assertEqual(GeneratedDocument.objects.filter(owner=self.user).count(), 1)

    def test_user_cannot_generate_from_another_account_template(self):
        self.client.force_authenticate(self.user)
        response = self.client.post(
            "/api/generate/",
            {"template_key": "private-template", "variables": {}},
            format="json",
        )

        self.assertEqual(response.status_code, status.HTTP_404_NOT_FOUND)

    def test_missing_required_fields_are_reported(self):
        self.client.force_authenticate(self.user)
        response = self.client.post(
            "/api/generate/",
            {
                "template_key": "introduction-letter",
                "variables": {"recipient_name": "شرکت نمونه"},
            },
            format="json",
        )

        self.assertEqual(response.status_code, status.HTTP_400_BAD_REQUEST)
        self.assertEqual(response.data["error"], "missing_required_fields")
        self.assertEqual(response.data["missing_fields"], ["subject"])

    def test_staff_can_view_all_documents_when_scope_all_is_requested(self):
        GeneratedDocument.objects.create(
            owner=self.user,
            created_by=self.user,
            template=self.template,
            variables={},
            status=GeneratedDocument.Status.SUCCESS,
        )
        GeneratedDocument.objects.create(
            owner=self.other_user,
            created_by=self.other_user,
            template=self.other_template,
            variables={},
            status=GeneratedDocument.Status.SUCCESS,
        )

        self.client.force_authenticate(self.staff)
        response = self.client.get("/api/documents/?scope=all")

        self.assertEqual(response.status_code, status.HTTP_200_OK)
        self.assertEqual(len(response.data["documents"]), 2)

    def test_service_api_key_can_list_templates_for_its_owner(self):
        _service_key, raw_key = ServiceAPIKey.build(
            owner=self.user,
            name="n8n production",
            allowed_sources=["n8n"],
        )
        self.client.credentials(HTTP_X_API_KEY=raw_key)

        response = self.client.get("/api/templates/")

        self.assertEqual(response.status_code, status.HTTP_200_OK)
        keys = [item["key"] for item in response.data["templates"]]
        self.assertEqual(keys, ["introduction-letter"])

    def test_service_api_key_can_generate_document_without_jwt(self):
        service_key, raw_key = ServiceAPIKey.build(
            owner=self.user,
            name="n8n production",
            allowed_sources=["n8n"],
        )
        self.client.credentials(HTTP_X_API_KEY=raw_key)

        response = self.client.post(
            "/api/generate/",
            {
                "template_key": "introduction-letter",
                "variables": {"recipient_name": "شرکت نمونه", "subject": "معرفی"},
                "idempotency_key": "api-key-execution-1",
                "source": "n8n",
            },
            format="json",
        )

        self.assertEqual(response.status_code, status.HTTP_201_CREATED)
        self.assertEqual(response.data["document"]["owner"]["username"], "n8n-account")
        service_key.refresh_from_db()
        self.assertIsNotNone(service_key.last_used_at)

    def test_service_api_key_rejects_disallowed_source(self):
        _service_key, raw_key = ServiceAPIKey.build(
            owner=self.user,
            name="n8n only",
            allowed_sources=["n8n"],
        )
        self.client.credentials(HTTP_X_API_KEY=raw_key)

        response = self.client.post(
            "/api/generate/",
            {
                "template_key": "introduction-letter",
                "variables": {"recipient_name": "شرکت نمونه", "subject": "معرفی"},
                "source": "api",
            },
            format="json",
        )

        self.assertEqual(response.status_code, status.HTTP_403_FORBIDDEN)
        self.assertEqual(response.data["error"], "source_not_allowed_for_api_key")

    def test_invalid_service_api_key_is_rejected(self):
        self.client.credentials(HTTP_X_API_KEY="n8n_invalid")
        response = self.client.get("/api/templates/")
        self.assertEqual(response.status_code, status.HTTP_401_UNAUTHORIZED)

